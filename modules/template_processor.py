import re
from docx import Document
from typing import List, Dict
import io
import zipfile


def _xml_replace_in_docx_bytes(docx_bytes: bytes, mapping: Dict[str, str]) -> bytes:
    """
    Perform raw XML replacements inside a DOCX (zip) file for any XML part
    under the `word/` folder. This helps replace placeholders that live
    inside shapes/textboxes or where runs split tokens.

    Args:
        docx_bytes: original docx as bytes
        mapping: dict of placeholder -> replacement

    Returns:
        Modified docx bytes
    """
    in_buf = io.BytesIO(docx_bytes)
    out_buf = io.BytesIO()

    with zipfile.ZipFile(in_buf, 'r') as zin:
        with zipfile.ZipFile(out_buf, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                # Only attempt textual replacements on word xml parts
                if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                    try:
                        text = data.decode('utf-8')
                    except Exception:
                        # If decoding fails, write the raw data back
                        zout.writestr(item, data)
                        continue

                    # For each key, replace variants: {{KEY}}, {{ KEY }}, [KEY], [ KEY ]
                    for key, value in mapping.items():
                        k = str(key).strip()
                        v = str(value)
                        # curly braces
                        text = re.sub(r"\{\{\s*" + re.escape(k) + r"\s*\}\}", v, text)
                        # bracket style
                        text = re.sub(r"\[\s*" + re.escape(k) + r"\s*\]", v, text)

                    zout.writestr(item, text.encode('utf-8'))
                else:
                    zout.writestr(item, data)

    return out_buf.getvalue()


def extract_placeholders(doc: Document) -> List[str]:
    """
    Extract all placeholders from a DOCX document.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        List of unique placeholder names (without {{ }})
    """
    placeholders = set()
    
    # Support both {{KEY}} and [KEY] placeholder styles
    placeholder_pattern = re.compile(r"{{(.*?)}}|\[(.*?)\]")

    for paragraph in doc.paragraphs:
        matches = placeholder_pattern.findall(paragraph.text)
        # matches is list of tuples because of alternation; pick the non-empty group
        for a, b in matches:
            name = (a or b)
            if name is not None:
                placeholders.add(name.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = placeholder_pattern.findall(paragraph.text)
                    for a, b in matches:
                        name = (a or b)
                        if name is not None:
                            placeholders.add(name.strip())

    # Also scan headers and footers
    try:
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                matches = placeholder_pattern.findall(paragraph.text)
                for a, b in matches:
                    name = (a or b)
                    if name is not None:
                        placeholders.add(name.strip())

            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = placeholder_pattern.findall(paragraph.text)
                            for a, b in matches:
                                name = (a or b)
                                if name is not None:
                                    placeholders.add(name.strip())

            footer = section.footer
            for paragraph in footer.paragraphs:
                matches = placeholder_pattern.findall(paragraph.text)
                for a, b in matches:
                    name = (a or b)
                    if name is not None:
                        placeholders.add(name.strip())

            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = placeholder_pattern.findall(paragraph.text)
                            for a, b in matches:
                                name = (a or b)
                                if name is not None:
                                    placeholders.add(name.strip())
    except Exception:
        pass
    
    return sorted(list(placeholders))


def find_unfilled_placeholders_from_bytes(doc_bytes: bytes) -> List[str]:
    """
    Load a DOCX from bytes and return any placeholders found (curly or bracket style).
    """
    doc = load_template(doc_bytes)
    return extract_placeholders(doc)


def apply_text_edits_to_docx_bytes(docx_bytes: bytes, edits: Dict[str, str]) -> bytes:
    """
    Apply suggested text edits (mapping from placeholder name or token to replacement text)
    directly into the DOCX bytes by performing XML-level replacements.
    Keys may be provided as 'INSURED_NAME' or as tokens like '[INSURED_NAME]' or '{{INSURED_NAME}}'.
    """
    normalized = {}
    for k, v in edits.items():
        key = str(k).strip()
        # if key looks like a bracketed or curly token, strip brackets
        if key.startswith('{{') and key.endswith('}}'):
            key = key[2:-2].strip()
        if key.startswith('[') and key.endswith(']'):
            key = key[1:-1].strip()
        normalized[key] = str(v)

    return _xml_replace_in_docx_bytes(docx_bytes, normalized)


def fill_docx_template(doc: Document, mapping: Dict[str, str]) -> Document:
    """
    Fill DOCX template by replacing placeholders with values.
    Handles cases where placeholders are split across multiple runs while preserving formatting.
    
    Args:
        doc: python-docx Document object
        mapping: Dictionary mapping placeholder names to values
        
    Returns:
        Modified Document object
    """
    def replace_paragraph_placeholders(paragraph, mapping):
        """Replace placeholders in a paragraph, preserving formatting where possible."""
        
        replaced_in_runs = False
        for run in paragraph.runs:
            run_changed = False
            run_text = run.text
            for key, value in mapping.items():
                k = str(key).strip()
                # Support both `{{KEY}}`, `{{ KEY }}` and `[KEY]`, `[ KEY ]` styles in runs
                placeholder_c = f"{{{{{k}}}}}"
                placeholder_c_spaced = f"{{{{ {k} }}}}"
                placeholder_b = f"[{k}]"
                placeholder_b_spaced = f"[ {k} ]"
                if placeholder_c in run_text or placeholder_c_spaced in run_text or placeholder_b in run_text or placeholder_b_spaced in run_text:
                    run_text = run_text.replace(placeholder_c, str(value))
                    run_text = run_text.replace(placeholder_c_spaced, str(value))
                    run_text = run_text.replace(placeholder_b, str(value))
                    run_text = run_text.replace(placeholder_b_spaced, str(value))
                    run_changed = True
            if run_changed:
                run.text = run_text
                replaced_in_runs = True
        
        if replaced_in_runs:
            return
        
        original_text = paragraph.text
        replaced_text = original_text
        for key, value in mapping.items():
            k = str(key).strip()
            placeholder_c = f"{{{{{k}}}}}"
            placeholder_c_spaced = f"{{{{ {k} }}}}"
            placeholder_b = f"[{k}]"
            placeholder_b_spaced = f"[ {k} ]"
            replaced_text = replaced_text.replace(placeholder_c, str(value))
            replaced_text = replaced_text.replace(placeholder_c_spaced, str(value))
            replaced_text = replaced_text.replace(placeholder_b, str(value))
            replaced_text = replaced_text.replace(placeholder_b_spaced, str(value))
        
        if replaced_text == original_text:
            return
        
        runs = list(paragraph.runs)
        if not runs:
            paragraph.add_run(replaced_text)
            return
        
        first_run_format = runs[0]
        
        for i in range(len(runs) - 1, -1, -1):
            runs[i]._element.getparent().remove(runs[i]._element)
        
        new_run = paragraph.add_run(replaced_text)
        new_run.bold = first_run_format.bold
        new_run.italic = first_run_format.italic
        new_run.underline = first_run_format.underline
        if first_run_format.style:
            new_run.style = first_run_format.style
        if first_run_format.font.size:
            new_run.font.size = first_run_format.font.size
        if first_run_format.font.name:
            new_run.font.name = first_run_format.font.name
        if first_run_format.font.color.rgb:
            new_run.font.color.rgb = first_run_format.font.color.rgb
    
    for paragraph in doc.paragraphs:
        replace_paragraph_placeholders(paragraph, mapping)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_placeholders(paragraph, mapping)

    # Replace in headers and footers for all sections
    try:
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                replace_paragraph_placeholders(paragraph, mapping)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_paragraph_placeholders(paragraph, mapping)

            footer = section.footer
            for paragraph in footer.paragraphs:
                replace_paragraph_placeholders(paragraph, mapping)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_paragraph_placeholders(paragraph, mapping)
    except Exception:
        # If headers/footers are not present or inaccessible, ignore and continue
        pass
    
    return doc


def load_template(file_bytes: bytes) -> Document:
    """
    Load a DOCX template from bytes.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        python-docx Document object
    """
    return Document(io.BytesIO(file_bytes))


def save_document(doc: Document, mapping: Dict[str, str] = None) -> bytes:
    """
    Save a DOCX document to bytes.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        DOCX file content as bytes
    """
    output = io.BytesIO()
    doc.save(output)
    raw = output.getvalue()

    # If a mapping is provided, run an XML-level replacement pass to
    # catch placeholders inside shapes/textboxes or other XML parts.
    if mapping:
        try:
            final = _xml_replace_in_docx_bytes(raw, mapping)
            return final
        except Exception:
            # On failure, return the original bytes
            return raw

    return raw
