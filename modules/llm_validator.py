import json
import io
from typing import Dict, Any, Optional
from google import genai
from docx import Document


def _extract_snippets_from_docx_bytes(doc_bytes: bytes, max_paragraphs: int = 8) -> str:
    try:
        doc = Document(io.BytesIO(doc_bytes))
        parts = []
        # headers
        for section in doc.sections:
            header_text = "\n".join(p.text for p in section.header.paragraphs if p.text)
            if header_text:
                parts.append("HEADER:\n" + header_text)
            break

        # first N paragraphs from body
        para_texts = [p.text for p in doc.paragraphs if p.text]
        parts.append("BODY_SNIPPETS:\n" + "\n".join(para_texts[:max_paragraphs]))

        # footer
        for section in doc.sections:
            footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text)
            if footer_text:
                parts.append("FOOTER:\n" + footer_text)
            break

        return "\n\n".join(parts)
    except Exception:
        return ""


def _strip_code_fences(text: str) -> str:
    t = text.strip()
    if t.startswith("```"):
        # remove leading ```json or ```
        idx = t.find('\n')
        if idx != -1:
            t = t[idx+1:]
        if t.endswith("```"):
            t = t[:-3]
    return t.strip()


def validate_filled_docx_bytes(doc_bytes: bytes, mapping: Dict[str, str], api_key: Optional[str]) -> Dict[str, Any]:
    """
    Validate a filled DOCX using Gemini LLM. Returns a dict with keys:
      - status: 'ok' | 'modify' | 'placeholders_left' | 'error'
      - issues: list of strings describing problems
      - suggested_mapping: optional dict of mapping overrides
      - suggested_text_edits: optional dict of paragraph replacements

    If the API fails or api_key is falsy, returns a conservative local-only result.
    """
    # Lightweight text extraction: we don't need entire docx content here —
    # include mapping and a small sample of text (mapping keys/values)
    try:
        client = genai.Client(api_key=api_key) if api_key else None
    except Exception:
        client = None

    payload = {"mapping": mapping}

    snippets = _extract_snippets_from_docx_bytes(doc_bytes, max_paragraphs=8)

    prompt = f"""
You are a document QA assistant. Given a document mapping (JSON) and a filled document snippets,
determine whether the document is ready for delivery. Respond with a VALID JSON object only.

Input mapping:
{json.dumps(mapping, indent=2)}

Document snippets (headers, first paragraphs, footer):
{snippets}

Respond with JSON with these keys:
- status: one of 'ok', 'modify', 'placeholders_left'
- issues: array of human strings describing problems
- suggested_mapping: optional JSON object with mapping updates (keys should be placeholder names, e.g. INSURED_NAME)
- suggested_text_edits: optional object mapping placeholder names or tokens to replacement text. Use placeholder names WITHOUT brackets, e.g. "INSURED_NAME": "Richard Daly".

Rules: If any placeholder tokens remain (strings like '{{...}}' or '[...]'), set status to 'placeholders_left' and include them in issues. If content reads awkward or needs polishing, set status to 'modify' and include suggested_text_edits. If everything is fine, set status to 'ok' and issues empty.
Return JSON only. No extra commentary.
"""

    if client is None:
        # No API access — fall back to simple placeholder check
        # The caller can run extract_placeholders on the filled doc to detect tokens.
        return {"status": "error", "issues": ["No API key available for LLM validation."], "suggested_mapping": {}, "suggested_text_edits": {}}

    try:
        response = client.models.generate_content(
            model="gemini-2.5",
            contents=prompt,
        )
        text = response.text or ""
        text = _strip_code_fences(text)
        # Attempt to parse JSON
        parsed = json.loads(text)
        # Normalize result
        result = {
            "status": parsed.get("status", "error"),
            "issues": parsed.get("issues", []),
            "suggested_mapping": parsed.get("suggested_mapping", {}),
            "suggested_text_edits": parsed.get("suggested_text_edits", {}),
        }
        return result
    except Exception as e:
        return {"status": "error", "issues": [f"LLM validation failed: {str(e)}"], "suggested_mapping": {}, "suggested_text_edits": {}}
