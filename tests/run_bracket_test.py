from docx import Document
from modules.template_processor import load_template, save_document, fill_docx_template, extract_placeholders

# Create a simple docx in memory with bracket placeholders
from io import BytesIO

doc = Document()
p = doc.add_paragraph()
p.add_run('Date of Loss: [DATE_LOSS]\n')
p2 = doc.add_paragraph()
p2.add_run('Insured: [ INSURED_NAME ]\n')
# Placeholder split across runs
p3 = doc.add_paragraph()
r1 = p3.add_run('Policy: [')
r2 = p3.add_run('POLICY_NUM')
r3 = p3.add_run(']')

buf = BytesIO()
doc.save(buf)
bytes_in = buf.getvalue()

template = load_template(bytes_in)
placeholders = extract_placeholders(template)
print('Extracted placeholders:', placeholders)

mapping = {'DATE_LOSS': '2025-11-01', 'INSURED_NAME': 'Bob', 'POLICY_NUM': 'P-12345'}
filled = fill_docx_template(template, mapping)

out_bytes = save_document(filled)
with open('tests/filled_bracket_test.docx', 'wb') as f:
    f.write(out_bytes)

print('Wrote tests/filled_bracket_test.docx')
from docx import Document as Doc
res_doc = Doc('tests/filled_bracket_test.docx')
texts = [p.text for p in res_doc.paragraphs]
print('Paragraphs after fill:')
for t in texts:
    print('-', t)

assert '2025-11-01' in '\n'.join(texts)
assert 'Bob' in '\n'.join(texts)
assert 'P-12345' in '\n'.join(texts)
print('Bracket replacement test PASSED.')
