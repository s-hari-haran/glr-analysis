from docx import Document
from modules.template_processor import load_template, save_document, fill_docx_template, extract_placeholders
from io import BytesIO

# Create docx with header/footer placeholders
doc = Document()
section = doc.sections[0]
header = section.header
header_p = header.paragraphs[0]
header_p.text = 'Header: [INSURED_NAME]'

p = doc.add_paragraph()
p.add_run('Date of Loss: [DATE_LOSS]')

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.text = 'Footer: [MORTGAGEE]'

buf = BytesIO()
doc.save(buf)
bytes_in = buf.getvalue()

template = load_template(bytes_in)
placeholders = extract_placeholders(template)
print('Extracted placeholders:', placeholders)

mapping = {
    'INSURED_NAME': 'Richard Daly',
    'DATE_LOSS': '9/28/2024',
    'MORTGAGEE': 'PennyMac'
}
filled = fill_docx_template(template, mapping)
out_bytes = save_document(filled)

with open('tests/filled_header_footer_test.docx', 'wb') as f:
    f.write(out_bytes)

# Verify
res = Document('tests/filled_header_footer_test.docx')
body_text = '\n'.join([p.text for p in res.paragraphs])
header_text = '\n'.join([p.text for p in res.sections[0].header.paragraphs])
footer_text = '\n'.join([p.text for p in res.sections[0].footer.paragraphs])
print('Body paragraphs:', body_text)
print('Header paragraphs:', header_text)
print('Footer paragraphs:', footer_text)
assert 'Richard Daly' in header_text
assert '9/28/2024' in body_text
assert 'PennyMac' in footer_text
print('Header/footer fill test PASSED')
