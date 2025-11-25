from docx import Document
from modules.template_processor import load_template, save_document, fill_docx_template, extract_placeholders

# Create a simple docx in memory with placeholders
doc = Document()
p = doc.add_paragraph()
p.add_run('Patient Name: {{NAME}}\n')
# Placeholder with spaces
p2 = doc.add_paragraph()
p2.add_run('Age: {{ AGE }}\n')
# Placeholder split across runs
p3 = doc.add_paragraph()
r1 = p3.add_run('{{')
r2 = p3.add_run('DOB')
r3 = p3.add_run('}}')

# Save to bytes
from io import BytesIO
buf = BytesIO()
doc.save(buf)
bytes_in = buf.getvalue()

# Load via template_processor
template = load_template(bytes_in)
placeholders = extract_placeholders(template)
print('Extracted placeholders:', placeholders)

mapping = {'NAME': 'Alice', 'AGE': '29', 'DOB': '1996-01-01'}
filled = fill_docx_template(template, mapping)

out_bytes = save_document(filled)

# Save to disk for manual inspection
with open('tests/filled_test.docx', 'wb') as f:
    f.write(out_bytes)

print('Wrote tests/filled_test.docx — open it to verify replacements.')

# Quick textual check: read paragraphs
from docx import Document as Doc
res_doc = Doc('tests/filled_test.docx')
texts = [p.text for p in res_doc.paragraphs]
print('Paragraphs after fill:')
for t in texts:
    print('-', t)

# Simple asserts
assert 'Alice' in '\n'.join(texts)
assert '29' in '\n'.join(texts)
assert '1996-01-01' in '\n'.join(texts)
print('All replacements detected in document — test PASSED.')
